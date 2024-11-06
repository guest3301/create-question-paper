from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class QuestionPaper:
    def __init__(self, template_path, output_path, global_font_size=12):
        self.doc = Document(template_path)
        self.output_path = output_path
        self.global_font_size = global_font_size
        self.set_page_margins()

    def replace_placeholders(self, class_name, subject, term, total_marks):
        for paragraph in self.doc.paragraphs:
            if '[CLASS]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[CLASS]', class_name)
                paragraph.runs[0].bold = True
            if '[SUBJECT]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[SUBJECT]', subject)
                paragraph.runs[0].bold = True
            if '[TERM]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[TERM]', term)
                paragraph.runs[0].bold = True
            if '[TMARKS]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[TMARKS]', total_marks)
                paragraph.runs[0].bold = True

    def set_page_margins(self):
        section = self.doc.sections[0]
        section.page_height = Inches(11.7)  # A4 height
        section.page_width = Inches(8.27)  # A4 width
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    def add_heading(self, text, font_size=None, bold=True):
        font_size = font_size or self.global_font_size
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_question(self, question_text, font_size=None, image_path=None, blank_lines=0, add_answer_line=False):
        font_size = font_size or self.global_font_size
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(question_text)
        run.font.size = Pt(font_size)
        if image_path:
            self.doc.add_paragraph().add_run().add_picture(image_path, width=Inches(1.0))
        for _ in range(blank_lines):
            self.doc.add_paragraph()
        if add_answer_line:
            self.add_answer_line()

    def add_answer_line(self, line_length=50):
        line = "_" * line_length
        paragraph = self.doc.add_paragraph(line)
        run = paragraph.add_run()
        run.font.size = Pt(self.global_font_size)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_section(self, heading_text, questions, marks, question_images=None, font_size=None, blank_lines=0, answer_lines=None):
        self.add_heading(f"{heading_text} ({marks} M)", font_size=font_size)
        if question_images and len(questions) != len(question_images):
            raise ValueError("The number of questions and images must match.")
        if answer_lines and len(questions) != len(answer_lines):
            raise ValueError("The number of questions and answer line indicators must match.")
        for idx, question in enumerate(questions):
            image_path = question_images[idx] if question_images else None
            add_answer_line = answer_lines[idx] if answer_lines else False
            self.add_question(question, font_size=font_size, image_path=image_path, blank_lines=blank_lines, add_answer_line=add_answer_line)

    def save(self):
        self.doc.save(self.output_path)
        print(f"Question paper saved as {self.output_path}")
'''
# Usage example
class_name = input("Enter the class name: ")
subject = input("Enter the subject: ")
term = input("Enter the term: ")
total_marks = input("Enter the total marks: ")
template_path = 'template.docx'   # Path to your document template
output_path = f'Question paper {class_name} ({subject}) [{term}].docx'
question_paper = QuestionPaper(template_path, output_path, global_font_size=12)
question_paper.replace_placeholders(class_name, subject, term, total_marks)
num_sections = int(input("Enter the number of sections: "))
for i in range(num_sections):
    print(f"\nSection {i + 1}:")
    heading_text = input("Enter the heading text (e.g., 'Q.1 Answer the following questions'): ")
    marks = input("Enter the marks for this section (e.g., '10 M'): ")
    blank_lines = int(input("Enter the number of blank lines after each question: "))
    questions = []
    question_images = []
    answer_lines = []
    num_questions = int(input("Enter the number of questions in this section: "))
    for j in range(num_questions):
        question_text = input(f"Enter question {j + 1} text: ")
        questions.append(question_text)
        image_path = input("Enter the image path for this question (or press Enter if none): ")
        question_images.append(image_path if image_path else None)
        add_answer_line = input("Does this question need an answer line? (yes/no): ").strip().lower() == "yes"
        answer_lines.append(add_answer_line)
    question_paper.add_section(
        heading_text=heading_text,
        questions=questions,
        marks=marks,
        question_images=question_images,
        font_size=12,
        blank_lines=blank_lines,
        answer_lines=answer_lines
    )

question_paper.save()
'''